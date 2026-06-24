from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_page_border(section):
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '24')
        border.set(qn('w:color'), 'C9A962')
        pgBorders.append(border)
    sectPr.append(pgBorders)

def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C9A962')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_code_block(doc, code, font_size=8):
    lines = code.split('\n')
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        # Add shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), 'F5F5F5')
        pPr.append(shd)

def create_styles(doc):
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level, size, color in [
        ('Heading 1', 22, '1A1A2E'),
        ('Heading 2', 16, '1A1A2E'),
        ('Heading 3', 13, 'C9A962'),
    ]:
        s = doc.styles[level]
        s.font.name = 'Calibri'
        s.font.bold = True
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor(*bytes.fromhex(color + '0' * (6 - len(color)) if len(color) < 6 else color[:6]))
        if ',' in color:
            parts = color.split(',')
            s.font.color.rgb = RGBColor(int(parts[0]), int(parts[1]), int(parts[2]))

def convert_table_lines(lines):
    rows = []
    for line in lines:
        if line.strip().startswith('|') and line.strip().endswith('|'):
            cells = [c.strip() for c in line.strip().split('|')[1:-1]]
            rows.append(cells)
    if len(rows) >= 2 and all(c == '' or all(ch == '-' for ch in c.replace(':', '')) for c in rows[1]):
        return [rows[0]] + rows[2:]
    return rows

def add_table(doc, rows):
    if len(rows) < 1:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            if i == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                set_cell_shading(cell, '1A1A2E')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif i % 2 == 0:
                set_cell_shading(cell, 'F9F9F9')
    doc.add_paragraph()

def add_list_item(doc, text, level=0, ordered=False, number=1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5 + level * 1)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    prefix = f"  {'  ' * level}• " if not ordered else f"  {'  ' * level}{number}. "
    run = p.add_run(prefix + text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'

def is_code_block_start(line):
    return line.strip().startswith('```')

def get_lang_from_code_start(line):
    return line.strip()[3:].strip()

def parse_and_generate(md_content, doc):
    create_styles(doc)
    
    # Title page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(100)
    run = p.add_run('ANGELOW STORE')
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = RGBColor(0xC9, 0xA9, 0x62)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Guía de Migración a Microservicios')
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run('Java 17+ · Spring Boot 3.x · MongoDB · Docker')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    run = p.add_run('Proyecto origen: Django 6.0.3 (Monolítico)')
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()
    add_page_border(doc.sections[0])

    # Table of Contents
    p = doc.add_paragraph()
    run = p.add_run('TABLA DE CONTENIDOS')
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    toc_items = [
        ('1.', 'Descripción General'),
        ('2.', 'Arquitectura de Microservicios Propuesta'),
        ('3.', 'Desglose de Microservicios', [
            ('3.1', 'Product Service (MS-01)'),
            ('3.2', 'Customer Service (MS-02)'),
            ('3.3', 'Order Service (MS-03)'),
            ('3.4', 'Inventory Service (MS-04)'),
            ('3.5', 'Payment Service (MS-05)'),
        ]),
        ('4.', 'Infraestructura Compartida'),
        ('5.', 'Paso a Paso de Implementación'),
        ('6.', 'Manejo de Estados (Workflows)'),
        ('7.', 'Comunicación entre Servicios'),
        ('8.', 'Manejo de Errores y Excepciones'),
        ('9.', 'Seguridad Distribuida'),
        ('10.', 'Configuraciones Clave'),
        ('11.', 'Resumen de Puertos'),
        ('12.', 'Checklist de Implementación'),
    ]
    for item in toc_items:
        num, title = item[0], item[1]
        add_list_item(doc, f'{num} {title}', level=0)
        if len(item) > 2:
            for sub in item[2]:
                add_list_item(doc, f'{sub[0]} {sub[1]}', level=1)

    doc.add_page_break()

    # Parse markdown
    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    code_buffer = []
    in_table = False
    table_buffer = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Handle code blocks
        if is_code_block_start(line):
            if in_code_block:
                add_code_block(doc, '\n'.join(code_buffer))
                code_buffer = []
                in_code_block = False
                i += 1
                continue
            else:
                in_code_block = True
                i += 1
                continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Handle tables
        if stripped.startswith('|') and stripped.endswith('|'):
            table_buffer.append(line)
            in_table = True
            i += 1
            continue
        else:
            if in_table and len(table_buffer) > 0:
                rows = convert_table_lines(table_buffer)
                if len(rows) >= 1:
                    add_table(doc, rows)
                table_buffer = []
                in_table = False
                if stripped == '':
                    i += 1
                    continue

        # Skip empty lines
        if stripped == '':
            i += 1
            continue

        # Horizontal rules
        if stripped == '---':
            add_horizontal_rule(doc)
            i += 1
            continue

        # Blockquotes
        if stripped.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            # Add left border
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '12')
            left.set(qn('w:space'), '8')
            left.set(qn('w:color'), 'C9A962')
            pBdr.append(left)
            pPr.append(pBdr)
            text = stripped[2:]
            # Bold
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Headers
        if stripped.startswith('##### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped[6:])
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            i += 1
            continue

        if stripped.startswith('#### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped[5:])
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            i += 1
            continue

        if stripped.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(stripped[4:])
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0xC9, 0xA9, 0x62)
            i += 1
            continue

        if stripped.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(stripped[3:])
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            i += 1
            continue

        if stripped.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(20)
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run(stripped[2:])
            run.bold = True
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            i += 1
            continue

        # Unordered list items
        if stripped.startswith('- ') or stripped.startswith('* '):
            level = 0
            text = stripped[2:]
            # Check for indentation-based nesting
            raw = line
            indent_match = re.match(r'^(\s+)', raw)
            if indent_match:
                level = len(indent_match.group(1)) // 2
            add_list_item(doc, text, level=level)
            i += 1
            continue

        # Checkbox list
        if stripped.startswith('- [') or stripped.startswith('* ['):
            level = 0
            raw = line
            indent_match = re.match(r'^(\s+)', raw)
            if indent_match:
                level = len(indent_match.group(1)) // 2
            checked = '[x]' in stripped.lower()
            text = stripped[stripped.index('] ') + 2:]
            prefix = '☑ ' if checked else '☐ '
            add_list_item(doc, prefix + text, level=level)
            i += 1
            continue

        # Ordered list
        if re.match(r'^\d+\.\s', stripped):
            num_match = re.match(r'^(\d+)\.\s(.*)', stripped)
            if num_match:
                add_list_item(doc, num_match.group(2), ordered=True, number=int(num_match.group(1)))
                i += 1
                continue

        # Bold text lines (like **text**)
        if re.match(r'^\*\*.*\*\*$', stripped):
            text = stripped.replace('**', '')
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            i += 1
            continue

        # Regular paragraph with possible formatting
        text = stripped
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
                run.font.size = Pt(10)
            else:
                run = p.add_run(part)
                run.font.size = Pt(10)

        i += 1

    # Last table check
    if in_table and len(table_buffer) > 0:
        rows = convert_table_lines(table_buffer)
        if len(rows) >= 1:
            add_table(doc, rows)

    # Add footer to last page
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    run = p.add_run('— Fin del documento —')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

def main():
    md_path = r'C:\Users\SENA\Downloads\angelow-store\ANGELOW-MICROSERVICES-PROMPT.md'
    docx_path = r'C:\Users\SENA\Downloads\angelow-store\docs\microservicio.docx'

    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    doc = Document()
    parse_and_generate(md_content, doc)
    doc.save(docx_path)
    print(f'Listo! Documento generado: {docx_path}')
    print(f'   Tamano: {__import__("os").path.getsize(docx_path) / 1024:.1f} KB')

if __name__ == '__main__':
    main()
